import io
import logging
import os
import platform
import subprocess
import tempfile
from datetime import datetime

logger = logging.getLogger(__name__)
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

LOGO_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "static", "Clip path group.svg"
)


def _add_horizontal_line(doc):
    """Add a thin horizontal line as a paragraph border."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from lxml import etree
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2B579A")


def _prepare_logo():
    """
    Convert the logo (SVG or PNG) to a flattened PNG on a white background
    so it renders with full brightness in Word.
    Returns an in-memory BytesIO stream.
    """
    from PIL import Image

    if LOGO_PATH.lower().endswith(".svg"):
        import cairosvg

        png_bytes = cairosvg.svg2png(url=LOGO_PATH, dpi=300)
        img = Image.open(io.BytesIO(png_bytes)).convert("RGBA")
    else:
        img = Image.open(LOGO_PATH).convert("RGBA")

    background = Image.new("RGBA", img.size, (255, 255, 255, 255))
    composite = Image.alpha_composite(background, img)
    composite = composite.convert("RGB")

    buf = io.BytesIO()
    composite.save(buf, format="PNG", dpi=(300, 300))
    buf.seek(0)
    return buf


def _add_logo_to_header(doc):
    """Add the CloudFuze logo to the document header, aligned top-right with spacing."""
    from docx.shared import Cm

    section = doc.sections[0]

    section.header_distance = Cm(1.0)
    section.top_margin = Cm(3.5)

    header = section.header
    header.is_linked_to_previous = False

    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_para.paragraph_format.space_after = Pt(12)

    run = header_para.add_run()
    if os.path.exists(LOGO_PATH):
        logo_stream = _prepare_logo()
        run.add_picture(logo_stream, width=Inches(1.6))


def generate_mom_document(
    meeting_title,
    meeting_date,
    meeting_time,
    tldr,
    action_items,
    decisions,
    transcript_text="",
):
    """
    Generate a professional Minutes of Meeting Word document.

    Returns the document as bytes (in-memory).
    """
    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # -- Logo in header (top-right) --
    _add_logo_to_header(doc)

    # -- Header Section --
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("MINUTES OF MEETING")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(43, 87, 154)

    sub_header = doc.add_paragraph()
    sub_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_header.add_run("CloudFuze, Inc.")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    _add_horizontal_line(doc)

    # -- Meeting Details Table --
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"

    details = [
        ("Meeting Title", meeting_title),
        ("Date", meeting_date),
        ("Time", meeting_time),
    ]

    for i, (label, value) in enumerate(details):
        row = table.rows[i]
        cell_label = row.cells[0]
        cell_value = row.cells[1]
        cell_label.text = label
        cell_value.text = value or "N/A"
        for paragraph in cell_label.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_paragraph()

    # -- Summary Section --
    heading = doc.add_heading("Summary", level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(43, 87, 154)

    if tldr:
        for point in tldr:
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(point)
    else:
        doc.add_paragraph("No summary provided.")

    # -- Action Items Section --
    heading = doc.add_heading("Action Items", level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(43, 87, 154)

    if action_items:
        action_table = doc.add_table(rows=1, cols=4)
        action_table.style = "Light Grid Accent 1"
        hdr = action_table.rows[0]
        hdr.cells[0].text = "#"
        hdr.cells[1].text = "Action Item"
        hdr.cells[2].text = "Assigned To"
        hdr.cells[3].text = "Due Date"
        for cell in hdr.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for idx, item in enumerate(action_items, 1):
            row = action_table.add_row()
            row.cells[0].text = str(idx)
            row.cells[1].text = item.get("description", "")
            row.cells[2].text = item.get("assigned_to", "")
            row.cells[3].text = item.get("due_date", "")
    else:
        doc.add_paragraph("No action items recorded.")

    # -- Decisions Section --
    heading = doc.add_heading("Decisions", level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(43, 87, 154)

    if decisions:
        for idx, decision in enumerate(decisions, 1):
            para = doc.add_paragraph(style="List Number")
            para.add_run(decision)
    else:
        doc.add_paragraph("No decisions recorded.")

    # -- Transcript Section (optional) --
    if transcript_text:
        doc.add_page_break()
        heading = doc.add_heading("Meeting Transcript", level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(43, 87, 154)

        para = doc.add_paragraph()
        run = para.add_run(transcript_text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)

    # -- Footer --
    _add_horizontal_line(doc)
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(
        "This document is confidential and intended solely for the recipients listed above.\n"
        f"Generated on {datetime.utcnow().strftime('%B %d, %Y at %I:%M %p')} UTC"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.italic = True

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def convert_docx_to_pdf(doc_bytes: bytes) -> bytes:
    """
    Convert a DOCX byte string to PDF.
    Uses LibreOffice on Linux, docx2pdf (Microsoft Word) on Windows.
    """
    with tempfile.TemporaryDirectory() as tmp_dir:
        docx_path = os.path.join(tmp_dir, "mom.docx")
        pdf_path = os.path.join(tmp_dir, "mom.pdf")

        with open(docx_path, "wb") as f:
            f.write(doc_bytes)

        if platform.system() == "Windows":
            try:
                from docx2pdf import convert as _convert
            except ImportError as exc:
                raise RuntimeError("docx2pdf is not installed. Run: pip install docx2pdf") from exc
            try:
                _convert(docx_path, pdf_path)
            except Exception as exc:
                logger.error("docx2pdf conversion failed: %s", exc)
                raise RuntimeError(f"PDF conversion failed: {exc}") from exc
        else:
            # LibreOffice headless — works on Linux/Mac without Microsoft Word
            libreoffice = _find_libreoffice()
            if not libreoffice:
                raise RuntimeError(
                    "LibreOffice is not installed. Run: apt-get install -y libreoffice"
                )
            result = subprocess.run(
                [libreoffice, "--headless", "--convert-to", "pdf", "--outdir", tmp_dir, docx_path],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode != 0:
                logger.error("LibreOffice conversion failed: %s", result.stderr)
                raise RuntimeError(f"PDF conversion failed: {result.stderr}")

        if not os.path.exists(pdf_path):
            raise RuntimeError("PDF conversion produced no output file.")

        with open(pdf_path, "rb") as f:
            return f.read()


def _find_libreoffice() -> str | None:
    """Return the LibreOffice executable path, or None if not found."""
    for candidate in ("libreoffice", "soffice"):
        result = subprocess.run(["which", candidate], capture_output=True, text=True)
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    return None
