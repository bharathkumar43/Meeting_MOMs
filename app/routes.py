import json
import logging
import os
import re
from datetime import datetime, timedelta
from functools import wraps

logger = logging.getLogger(__name__)

from flask import (
    Blueprint, render_template, redirect, url_for,
    request, session, flash, send_file, jsonify,
)
import io
from docx import Document
from werkzeug.utils import secure_filename

from app.auth import get_auth_url, process_auth_callback, get_token, get_app_token, logout as auth_logout
from app.graph_client import GraphClient
from app.zoom_auth import get_zoom_access_token
from app.zoom_client import ZoomClient
from app.meeting_filter import (
    filter_customer_meetings, filter_by_subject,
    parse_vtt_transcript, transcript_to_readable,
)
from app.doc_generator import generate_mom_document, convert_docx_to_pdf
from app.email_sender import send_mom_email
from app.mom_generator import generate_mom_from_transcript
from app.activity_tracker import (
    record_login, record_meeting_access, record_mom_sent,
    get_all_users, get_user_stats, get_pending_moms, get_sent_moms,
    get_managers, get_non_managers,
    get_audit_rows, get_audit_totals,
    get_all_user_dashboard_stats,
)
from app.audit_report_email import send_daily_audit_report
from config import Config

main_bp = Blueprint("main", __name__)

TRANSCRIPT_UPLOAD_MAX_BYTES = 10 * 1024 * 1024
ALLOWED_TRANSCRIPT_SUFFIXES = {".txt", ".vtt", ".docx"}


def _decode_text_file(raw: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _text_from_uploaded_vtt(raw: bytes) -> str:
    text = _decode_text_file(raw)
    entries = parse_vtt_transcript(text)
    if entries:
        return transcript_to_readable(entries).strip()
    return text.strip()


def _text_from_uploaded_docx(raw: bytes) -> str:
    doc = Document(io.BytesIO(raw))
    parts = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)
    return "\n".join(parts).strip()


def _normalize_subject(s):
    """Normalize a meeting subject for comparison.
    Teams strips special chars like | from recording filenames, so we must
    remove them and collapse whitespace for both sides to match."""
    s = s.lower().strip()
    s = re.sub(r'[|/\\:*?"<>]', ' ', s)
    s = re.sub(r'\s+', ' ', s)
    return s


def _match_recording(meeting, all_recordings):
    """Check if a calendar event has a matching recording file in OneDrive."""
    if not all_recordings:
        return False
    meeting_start = meeting.get("start", {}).get("dateTime", "")
    if not meeting_start:
        return False
    meeting_date = meeting_start[:10].replace("-", "")
    norm_subject = _normalize_subject(meeting.get("subject", ""))
    for rec in all_recordings:
        if rec["date"] != meeting_date:
            continue
        norm_rec = _normalize_subject(rec["subject"])
        if norm_rec == norm_subject or norm_rec in norm_subject or norm_subject in norm_rec:
            return True
    return False


def _app_user_ids_for_transcript_lookup(app_client, meeting, signed_in_email="", viewed_user_id=None):
    """Ordered Graph user IDs to try for app-only onlineMeeting + transcript lookup."""
    ids = []
    seen = set()

    def add_uid(uid):
        if uid and uid not in seen:
            seen.add(uid)
            ids.append(uid)

    if viewed_user_id:
        add_uid(viewed_user_id)

    organizer_email = (
        meeting.get("organizer", {}).get("emailAddress", {}).get("address", "") or ""
    )
    if organizer_email:
        add_uid(app_client._resolve_user_id_safe(organizer_email))

    if signed_in_email and signed_in_email.lower() != organizer_email.lower():
        add_uid(app_client._resolve_user_id_safe(signed_in_email))

    return ids


def _graph_transcript_exists_app(app_client, join_url, user_ids):
    """True if any listed user can resolve join_url to an online meeting with transcripts."""
    if not join_url or not app_client or not user_ids:
        return False
    for uid in user_ids:
        if not uid:
            continue
        try:
            online_meeting = app_client.get_online_meeting_for_user(uid, join_url)
            if not online_meeting:
                continue
            meeting_id = online_meeting["id"]
            transcripts = app_client.list_transcripts(meeting_id, user_id=uid)
            if transcripts:
                return True
        except Exception as e:
            logger.debug("App transcript lookup failed for uid=%s: %s", uid, e)
            continue
    return False


def meeting_has_transcript_signal(
    meeting,
    all_recordings,
    delegated_client,
    app_client,
    signed_in_email="",
    viewed_user_id=None,
):
    """
    True if we expect transcript content to be available: OneDrive recording match
    OR Graph lists at least one transcript (delegated /me path, then app fallback).
    """
    if _match_recording(meeting, all_recordings):
        return True

    join_url = (meeting.get("onlineMeeting") or {}).get("joinUrl", "") or ""
    if not join_url:
        return False

    if delegated_client:
        try:
            if delegated_client.check_transcript_exists(join_url):
                return True
        except Exception as e:
            logger.debug("Delegated transcript check failed: %s", e)

    if app_client:
        ids = _app_user_ids_for_transcript_lookup(
            app_client, meeting, signed_in_email=signed_in_email, viewed_user_id=viewed_user_id
        )
        if _graph_transcript_exists_app(app_client, join_url, ids):
            return True

    return False


def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = get_token()
        if not token:
            return redirect(url_for("main.login_page"))
        return f(*args, **kwargs)
    return decorated


def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = get_token()
        if not token:
            return redirect(url_for("main.login_page"))
        user = session.get("user", {})
        email = (user.get("preferred_username") or user.get("upn") or "").lower()
        if email not in Config.ADMIN_EMAILS:
            flash("Access denied. Admin privileges required.", "danger")
            return redirect(url_for("main.dashboard"))
        return f(*args, **kwargs)
    return decorated


# ── Auth Routes ──────────────────────────────────────────────

@main_bp.route("/")
def index():
    if get_token():
        return redirect(url_for("main.dashboard"))
    return render_template("login.html")


@main_bp.route("/login")
def login():
    auth_url = get_auth_url()
    return redirect(auth_url)


@main_bp.route("/login-page")
def login_page():
    return render_template("login.html")


@main_bp.route("/auth/callback")
def auth_callback():
    result = process_auth_callback()
    if result:
        user = session.get("user", {})
        email = (user.get("preferred_username") or user.get("upn") or "").lower()
        name = user.get("name", "")
        if email:
            record_login(email, name)
        flash("Signed in successfully.", "success")
        return redirect(url_for("main.dashboard"))
    flash("Authentication failed. Please try again.", "danger")
    return redirect(url_for("main.login_page"))


@main_bp.route("/logout")
def logout():
    auth_logout()
    flash("You have been signed out.", "info")
    return redirect(url_for("main.index"))


@main_bp.route("/admin/grant-consent")
@admin_required
def admin_grant_consent():
    """Redirect admin to Azure AD admin consent page."""
    consent_url = (
        f"https://login.microsoftonline.com/{Config.AZURE_TENANT_ID}"
        f"/adminconsent?client_id={Config.AZURE_CLIENT_ID}"
        f"&redirect_uri={Config.REDIRECT_URI}"
    )
    return redirect(consent_url)


# ── Dashboard ────────────────────────────────────────────────

@main_bp.route("/dashboard")
@login_required
def dashboard():
    start_date = request.args.get("start_date", "")
    end_date = request.args.get("end_date", "")
    keyword = request.args.get("keyword", "")

    meetings = None
    explicitly_searched = bool(request.args.get("start_date") or request.args.get("end_date"))

    if not start_date:
        start_date = (datetime.now() - timedelta(days=7)).strftime("%Y-%m-%d")
    if not end_date:
        end_date = datetime.now().strftime("%Y-%m-%d")

    if explicitly_searched:
        try:
            token = get_token()
            client = GraphClient(token)
            logger.info("Dashboard: fetching events from %s to %s", start_date, end_date)
            events = client.list_calendar_events(start_date, end_date)
            logger.info("Dashboard: got %d online meeting events", len(events))

            meetings = filter_customer_meetings(events, Config.ORG_DOMAIN)
            logger.info("Dashboard: %d customer meetings after filter", len(meetings))

            if keyword:
                keywords = [k.strip() for k in keyword.split(",")]
                meetings = filter_by_subject(meetings, keywords)

            app_token = get_app_token()
            app_client = GraphClient(app_token) if app_token else None

            all_recordings = []
            if app_client:
                user_profile = client.get_user_profile()
                dashboard_user_id = user_profile.get("id", "")
                if dashboard_user_id:
                    all_recordings = app_client.build_recording_lookup(
                        meetings, dashboard_user_id, Config.ORG_DOMAIN
                    )

            for meeting in meetings:
                meeting["has_transcript"] = _match_recording(meeting, all_recordings)

        except Exception as e:
            flash(f"Error fetching meetings: {str(e)}", "danger")
            meetings = []

    return render_template(
        "dashboard.html",
        meetings=meetings,
        start_date=start_date,
        end_date=end_date,
        keyword=keyword,
    )


# ── Transcript View ──────────────────────────────────────────

@main_bp.route("/transcript")
@login_required
def transcript():
    event_id = request.args.get("event_id", "")
    join_url = request.args.get("join_url", "")

    token = get_token()
    client = GraphClient(token)

    # Fetch the calendar event directly by ID (avoids loading 90 days of events)
    try:
        meeting = client.get_event(event_id)
    except Exception:
        meeting = None

    if not meeting:
        flash("Meeting not found.", "warning")
        return redirect(url_for("main.dashboard"))

    # Build attendees list
    attendees = []
    for att in meeting.get("attendees", []):
        attendees.append({
            "name": att.get("emailAddress", {}).get("name", ""),
            "email": att.get("emailAddress", {}).get("address", ""),
        })

    transcript_text = ""
    error = None  # kept for template compatibility; transcript fetch warnings are not shown
    organizer_email = (
        meeting.get("organizer", {}).get("emailAddress", {}).get("address", "")
    )

    # Use the join URL from the calendar event directly (avoids URL encoding issues)
    meeting_join_url = (
        meeting.get("onlineMeeting", {}).get("joinUrl", "")
    )
    effective_join_url = meeting_join_url or join_url

    logger.info(
        "Transcript request: subject='%s', organizer='%s', "
        "join_url_param=%d chars, meeting_join_url=%d chars",
        meeting.get("subject", ""),
        organizer_email,
        len(join_url),
        len(meeting_join_url),
    )

    # Step 1: Try delegated token (works when user has OnlineMeetings.Read consent)
    try:
        online_meeting = client.get_online_meeting_by_join_url(effective_join_url)
        if online_meeting:
            meeting_id = online_meeting["id"]
            logger.info("Step 1 OK: Found online meeting via delegated token")
            transcripts = client.list_transcripts(meeting_id)
            logger.info("Step 1: %d transcript(s) found", len(transcripts))
            if transcripts:
                transcript_id = transcripts[0]["id"]
                vtt_content = client.get_transcript_content(meeting_id, transcript_id)
                if vtt_content:
                    entries = parse_vtt_transcript(vtt_content)
                    transcript_text = transcript_to_readable(entries)
                    logger.info("Step 1 SUCCESS: loaded %d chars", len(transcript_text))
        else:
            logger.info("Step 1: No online meeting found via delegated token")
    except Exception as e:
        logger.warning("Step 1 FAILED: %s", e)

    # Step 2: If delegated didn't work, try app token
    if not transcript_text:
        app_token = get_app_token()
        if app_token and effective_join_url:
            app_client = GraphClient(app_token)

            user_data = session.get("user", {})
            signed_in_email = (
                user_data.get("preferred_username") or user_data.get("upn") or ""
            )
            targets = []
            if organizer_email:
                targets.append(("organizer", organizer_email))
            if signed_in_email and signed_in_email.lower() != (organizer_email or "").lower():
                targets.append(("signed-in user", signed_in_email))

            for label, target_email in targets:
                if transcript_text:
                    break
                try:
                    uid = app_client._resolve_user_id_safe(target_email)
                    if not uid:
                        logger.info("Step 2: Cannot resolve %s (%s)", label, target_email)
                        continue
                    logger.info("Step 2: Trying via %s (%s)", label, target_email)
                    online_meeting = app_client.get_online_meeting_for_user(
                        uid, effective_join_url
                    )
                    if not online_meeting:
                        logger.info("Step 2: No meeting found via %s", label)
                        continue
                    meeting_id = online_meeting["id"]
                    logger.info("Step 2: Found meeting via %s", label)
                    transcripts = app_client.list_transcripts(meeting_id, user_id=uid)
                    logger.info("Step 2: %d transcript(s) via %s", len(transcripts), label)
                    if transcripts:
                        transcript_id = transcripts[0]["id"]
                        vtt_content = app_client.get_transcript_content(
                            meeting_id, transcript_id, user_id=uid
                        )
                        if vtt_content:
                            entries = parse_vtt_transcript(vtt_content)
                            transcript_text = transcript_to_readable(entries)
                            logger.info("Step 2 SUCCESS via %s: %d chars", label, len(transcript_text))
                except Exception as e:
                    logger.warning("Step 2 FAILED via %s: %s", label, e)

    user = session.get("user", {})
    email = (user.get("preferred_username") or user.get("upn") or "").lower()
    if email:
        record_meeting_access(
            email,
            meeting.get("subject", ""),
            meeting.get("start", {}).get("dateTime", "")[:10],
        )

    return render_template(
        "transcript.html",
        meeting=meeting,
        transcript_text=transcript_text,
        attendees_json=json.dumps(attendees),
        error=error,
    )


@main_bp.route("/transcript/parse-file", methods=["POST"])
@login_required
def parse_transcript_file():
    """Extract plain text from an uploaded .txt, .vtt, or .docx transcript file."""
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded."}), 400
    upload = request.files["file"]
    if not upload or not upload.filename:
        return jsonify({"error": "No file selected."}), 400

    safe_name = secure_filename(upload.filename)
    _, ext = os.path.splitext(safe_name.lower())
    if ext not in ALLOWED_TRANSCRIPT_SUFFIXES:
        allowed = ", ".join(sorted(ALLOWED_TRANSCRIPT_SUFFIXES))
        return jsonify({"error": f"Unsupported type. Allowed: {allowed}"}), 400

    raw = upload.read(TRANSCRIPT_UPLOAD_MAX_BYTES + 1)
    if len(raw) > TRANSCRIPT_UPLOAD_MAX_BYTES:
        return jsonify({"error": "File is too large (maximum 10 MB)."}), 413

    try:
        if ext == ".txt":
            text = _decode_text_file(raw).strip()
        elif ext == ".vtt":
            text = _text_from_uploaded_vtt(raw)
        else:
            text = _text_from_uploaded_docx(raw)
    except Exception as e:
        logger.exception("parse_transcript_file failed")
        return jsonify({"error": f"Could not read this file: {str(e)}"}), 400

    if not text or not text.strip():
        return jsonify({"error": "No text could be extracted from this file."}), 400

    return jsonify({"text": text})


# ── MOM Builder ──────────────────────────────────────────────

@main_bp.route("/mom-builder", methods=["POST"])
@login_required
def mom_builder():
    meeting_subject = request.form.get("meeting_subject", "")
    meeting_date = request.form.get("meeting_date", "")
    meeting_time = request.form.get("meeting_time", "")
    attendees_json = request.form.get("attendees_json", "[]")
    transcript = request.form.get("transcript", "")

    mom, gen_error = generate_mom_from_transcript(transcript, meeting_subject)
    if gen_error:
        flash(f"AI generation: {gen_error}", "warning")

    return render_template(
        "mom_builder.html",
        meeting_subject=meeting_subject,
        meeting_date=meeting_date,
        meeting_time=meeting_time,
        attendees_json=attendees_json,
        transcript=transcript,
        ai_tldr=mom["tldr"],
        ai_action_items=mom["action_items"],
        ai_decisions=mom["decisions"],
    )


# ── Send Page ────────────────────────────────────────────────

@main_bp.route("/send", methods=["POST"])
@login_required
def send_page():
    meeting_subject = request.form.get("meeting_subject", "")
    meeting_date = request.form.get("meeting_date", "")
    meeting_time = request.form.get("meeting_time", "")
    attendees_json = request.form.get("attendees_json", "[]")
    transcript = request.form.get("transcript", "")

    title = request.form.get("title", meeting_subject)
    tldr = [p for p in request.form.getlist("tldr") if p.strip()]

    action_descs = request.form.getlist("action_desc")
    action_assignees = request.form.getlist("action_assignee")
    action_dues = request.form.getlist("action_due")
    action_items = []
    for desc, assignee, due in zip(action_descs, action_assignees, action_dues):
        if desc.strip():
            action_items.append({
                "description": desc,
                "assigned_to": assignee,
                "due_date": due,
            })

    decisions = request.form.getlist("decisions")
    decisions = [d for d in decisions if d.strip()]

    include_transcript = request.form.get("include_transcript") == "on"

    attendees = json.loads(attendees_json)

    doc_bytes = generate_mom_document(
        meeting_title=title,
        meeting_date=meeting_date,
        meeting_time=meeting_time,
        tldr=tldr,
        action_items=action_items,
        decisions=decisions,
        transcript_text=transcript if include_transcript else "",
    )

    # Extract external attendee emails for pre-filling
    org_domain = Config.ORG_DOMAIN.lower()
    external_attendees = [
        a for a in attendees
        if a.get("email") and not a["email"].lower().endswith(f"@{org_domain}")
    ]
    external_emails = [a["email"] for a in external_attendees]
    primary_email = external_emails[0] if external_emails else ""

    # Resolve greeting name from the local part of the primary external email
    # e.g. john.doe@example.com → "John"
    if primary_email:
        local_part = primary_email.split("@")[0]
        first_token = re.split(r"[._\-]", local_part)[0]
        greeting_name = first_token.capitalize() if first_token else "Customer"
    else:
        greeting_name = "Customer"

    # Build CC list: remaining external attendees + default team addresses
    cc_parts = external_emails[1:]
    for default_addr in [Config.MOM_DEFAULT_DL_EMAIL, Config.MOM_DEFAULT_CC_EMAIL]:
        if default_addr and default_addr not in cc_parts:
            cc_parts.append(default_addr)
    prefill_cc = ", ".join(cc_parts)

    has_default_cc = bool(Config.MOM_DEFAULT_DL_EMAIL or Config.MOM_DEFAULT_CC_EMAIL)

    session["mom_doc"] = {
        "bytes_hex": doc_bytes.hex(),
        "meeting_subject": title,
        "meeting_date": meeting_date,
        "meeting_time": meeting_time,
        "attendees": attendees,
    }

    return render_template(
        "send.html",
        meeting_subject=title,
        meeting_date=meeting_date,
        meeting_time=meeting_time,
        action_count=len(action_items),
        decision_count=len(decisions),
        sent_success=False,
        prefill_email=primary_email,
        prefill_cc=prefill_cc,
        has_default_cc=has_default_cc,
        greeting_name=greeting_name,
    )


# ── Download MOM ─────────────────────────────────────────────

@main_bp.route("/download-mom")
@login_required
def download_mom():
    mom_data = session.get("mom_doc")
    if not mom_data:
        flash("No document available to download. Please generate a MOM first.", "warning")
        return redirect(url_for("main.dashboard"))

    doc_bytes = bytes.fromhex(mom_data["bytes_hex"])
    safe_title = "".join(
        c if c.isalnum() or c in " -_" else "_"
        for c in mom_data["meeting_subject"]
    )
    filename = f"MOM_{safe_title}_{mom_data['meeting_date']}.docx"

    return send_file(
        io.BytesIO(doc_bytes),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


# ── Send Email ───────────────────────────────────────────────

@main_bp.route("/send-email", methods=["POST"])
@login_required
def send_email():
    mom_data = session.get("mom_doc")
    if not mom_data:
        flash("No document available. Please generate a MOM first.", "warning")
        return redirect(url_for("main.dashboard"))

    to_email = request.form.get("to_email", "").strip()
    cc_emails = request.form.get("cc_emails", "").strip()
    format_choice = request.form.get("format_choice", "docx")  # "docx" | "pdf"
    greeting_name = request.form.get("greeting_name", "").strip() or "Customer"

    to_list = [e.strip() for e in to_email.split(",") if e.strip()]
    cc_list = [e.strip() for e in cc_emails.split(",") if e.strip()]

    if not to_list:
        flash("Please provide at least one customer email address.", "warning")
        return redirect(url_for("main.send_page"))

    try:
        token = get_token()
        doc_bytes = bytes.fromhex(mom_data["bytes_hex"])

        safe_title = "".join(
            c if c.isalnum() or c in " -_" else "_"
            for c in mom_data["meeting_subject"]
        )
        base_name = f"MOM_{safe_title}_{mom_data['meeting_date']}"

        # Build MOM attachment based on format selection
        attachments = []
        if format_choice == "pdf":
            pdf_bytes = convert_docx_to_pdf(doc_bytes)
            attachments.append({
                "bytes": pdf_bytes,
                "filename": f"{base_name}.pdf",
                "content_type": "application/pdf",
            })
        else:
            attachments.append({
                "bytes": doc_bytes,
                "filename": f"{base_name}.docx",
                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            })

        # Append any extra file the user uploaded
        extra_file = request.files.get("extra_attachment")
        if extra_file and extra_file.filename:
            extra_bytes = extra_file.read()
            max_bytes = 20 * 1024 * 1024  # 20 MB
            if len(extra_bytes) > max_bytes:
                flash("Extra attachment exceeds the 20 MB size limit.", "warning")
                return redirect(url_for("main.send_page"))
            safe_extra_name = secure_filename(extra_file.filename)
            attachments.append({
                "bytes": extra_bytes,
                "filename": safe_extra_name,
                "content_type": extra_file.content_type or "application/octet-stream",
            })

        filename, sent_list = send_mom_email(
            access_token=token,
            to_emails=to_list,
            cc_emails=cc_list,
            meeting_title=mom_data["meeting_subject"],
            meeting_date=mom_data["meeting_date"],
            attachments=attachments,
            greeting_name=greeting_name,
        )

        user = session.get("user", {})
        email = (user.get("preferred_username") or user.get("upn") or "").lower()
        if email:
            for recipient in sent_list:
                record_mom_sent(
                    email=email,
                    subject=mom_data["meeting_subject"],
                    meeting_date=mom_data["meeting_date"],
                    sent_to=recipient.strip(),
                )

        sent_display = ", ".join(sent_list)
        flash(f"MOM sent successfully to {sent_display}!", "success")

        return render_template(
            "send.html",
            meeting_subject=mom_data["meeting_subject"],
            meeting_date=mom_data["meeting_date"],
            meeting_time=mom_data["meeting_time"],
            action_count=0,
            decision_count=0,
            sent_success=True,
            sent_to=sent_display,
        )

    except Exception as e:
        flash(f"Failed to send email: {str(e)}", "danger")
        return render_template(
            "send.html",
            meeting_subject=mom_data["meeting_subject"],
            meeting_date=mom_data["meeting_date"],
            meeting_time=mom_data["meeting_time"],
            action_count=0,
            decision_count=0,
            sent_success=False,
        )


# ── Admin Dashboard ───────────────────────────────────────────

@main_bp.route("/admin")
@admin_required
def admin_dashboard():
    users = get_all_users()
    managers = get_managers()
    non_managers = get_non_managers()
    total_users, total_meetings, total_sent = get_user_stats()
    pending_moms = get_pending_moms()
    sent_moms = get_sent_moms()
    user_dashboard_stats = get_all_user_dashboard_stats()

    manager_emails = Config.MANAGER_EMAILS
    manager_pending = [p for p in pending_moms if p.user_email in manager_emails]
    manager_sent = [s for s in sent_moms if s.user_email in manager_emails]

    audit_days = max(1, Config.AUDIT_REPORT_DAYS)
    audit_rows = get_audit_rows(days=audit_days)
    audit_totals = get_audit_totals(audit_rows)
    audit_can_send = bool(
        Config.AUDIT_SENDER_MAILBOX and Config.AUDIT_REPORT_RECIPIENTS
    )
    audit_sender_display = Config.AUDIT_SENDER_MAILBOX or "—"
    audit_recipients_display = ", ".join(Config.AUDIT_REPORT_RECIPIENTS) or "—"

    return render_template(
        "admin.html",
        users=users,
        managers=managers,
        non_managers=non_managers,
        total_users=total_users,
        total_meetings=total_meetings,
        total_sent=total_sent,
        total_pending=len(pending_moms),
        pending_moms=pending_moms,
        sent_moms=sent_moms,
        manager_pending=manager_pending,
        manager_sent=manager_sent,
        audit_rows=audit_rows,
        audit_totals=audit_totals,
        audit_days=audit_days,
        audit_can_send=audit_can_send,
        audit_sender_display=audit_sender_display,
        audit_recipients_display=audit_recipients_display,
        user_dashboard_stats=user_dashboard_stats,
    )


@main_bp.route("/admin/send-audit-report", methods=["POST"])
@admin_required
def admin_send_audit_report():
    """Send the rolling audit email now (Graph app-only)."""
    ok, msg = send_daily_audit_report(force=True)
    if ok:
        flash(msg, "success")
    else:
        flash(msg, "danger")
    return redirect(url_for("main.admin_dashboard") + "#audit-report")


# ── Admin: Per-User Meetings View ────────────────────────────

@main_bp.route("/admin/user-meetings/<path:user_email>")
@admin_required
def admin_user_meetings(user_email):
    today = datetime.now().strftime("%Y-%m-%d")
    default_start = (datetime.now() - timedelta(days=7)).strftime("%Y-%m-%d")
    start_date = request.args.get("start_date") or default_start
    end_date = request.args.get("end_date") or today

    meetings = []
    user_display_name = user_email
    error = None

    try:
        app_token = get_app_token()
        if not app_token:
            error = (
                "Application token not available. "
                "Ensure Application permissions (Calendars.Read, OnlineMeetingTranscript.Read.All) "
                "are configured in Azure AD with admin consent."
            )
        else:
            client = GraphClient(app_token)

            user_info = client.get_user_id(user_email)
            user_id = user_info["id"]
            user_display_name = user_info.get("displayName", user_email)

            events = client.list_user_calendar_events(user_id, start_date, end_date)
            meetings = filter_customer_meetings(events, Config.ORG_DOMAIN)

            all_recordings = client.build_recording_lookup(
                meetings, user_id, Config.ORG_DOMAIN
            )

            for meeting in meetings:
                meeting["has_transcript"] = _match_recording(meeting, all_recordings)

    except Exception as e:
        error = f"Error fetching meetings for {user_email}: {str(e)}"
        meetings = []

    from app.models import MOMSent, User as UserModel
    sent_keys = set()
    db_user = UserModel.query.filter_by(email=user_email.lower()).first()
    if db_user:
        sent_records = db_user.sent_moms.all()
        for s in sent_records:
            sent_keys.add(f"{s.subject}|{s.meeting_date}")

    return render_template(
        "admin_user_meetings.html",
        user_email=user_email,
        user_display_name=user_display_name,
        meetings=meetings,
        sent_keys=sent_keys,
        start_date=start_date,
        end_date=end_date,
        error=error,
    )


# ── Zoom Dashboard ────────────────────────────────────────────

@main_bp.route("/zoom/dashboard")
@login_required
def zoom_dashboard():
    start_date = request.args.get("start_date", "")
    end_date = request.args.get("end_date", "")
    keyword = request.args.get("keyword", "")

    explicitly_searched = bool(request.args.get("start_date") or request.args.get("end_date"))

    if not start_date:
        start_date = (datetime.now() - timedelta(days=7)).strftime("%Y-%m-%d")
    if not end_date:
        end_date = datetime.now().strftime("%Y-%m-%d")

    user_data = session.get("user", {})
    user_email = (user_data.get("preferred_username") or user_data.get("upn") or "").lower()

    zoom_token = get_zoom_access_token()
    zoom_configured = bool(zoom_token)

    meetings = None
    if explicitly_searched and zoom_token and user_email:
        try:
            zoom = ZoomClient(zoom_token)
            recordings = zoom.list_user_recordings(user_email, start_date, end_date)
            meetings = [ZoomClient.normalize_recording(r) for r in recordings]

            if keyword:
                keywords = [k.strip() for k in keyword.split(",")]
                meetings = filter_by_subject(meetings, keywords)

        except Exception as e:
            flash(f"Error fetching Zoom recordings: {str(e)}", "danger")
            meetings = []

    return render_template(
        "zoom_dashboard.html",
        meetings=meetings,
        start_date=start_date,
        end_date=end_date,
        keyword=keyword,
        zoom_configured=zoom_configured,
    )


# ── Zoom Transcript View ──────────────────────────────────────

@main_bp.route("/zoom/transcript")
@login_required
def zoom_transcript():
    meeting_id = request.args.get("meeting_id", "")
    if not meeting_id:
        flash("No meeting ID provided.", "warning")
        return redirect(url_for("main.zoom_dashboard"))

    zoom_token = get_zoom_access_token()
    if not zoom_token:
        flash("Zoom integration is not configured. Add Zoom credentials to your .env file.", "danger")
        return redirect(url_for("main.zoom_dashboard"))

    zoom = ZoomClient(zoom_token)

    recording = zoom.get_meeting_recordings(meeting_id)
    if not recording:
        flash("Could not retrieve Zoom meeting recording.", "warning")
        return redirect(url_for("main.zoom_dashboard"))

    meeting = ZoomClient.normalize_recording(recording)

    transcript_text = ""
    if meeting.get("zoom_transcript_url"):
        try:
            vtt_content = zoom.get_transcript_content(meeting["zoom_transcript_url"])
            if vtt_content:
                entries = parse_vtt_transcript(vtt_content)
                transcript_text = transcript_to_readable(entries)
                logger.info("Zoom transcript loaded: %d chars for meeting %s", len(transcript_text), meeting_id)
        except Exception as e:
            logger.warning("Zoom transcript fetch failed for %s: %s", meeting_id, e)

    attendees = []
    external_attendees = []
    try:
        participants = zoom.get_meeting_participants(meeting_id)
        org_domain = Config.ORG_DOMAIN.lower()
        for p in participants:
            email = (p.get("user_email") or "").strip()
            name = (p.get("name") or p.get("user_name") or "").strip()
            if email:
                attendees.append({"name": name, "email": email})
                if not email.lower().endswith(f"@{org_domain}"):
                    external_attendees.append({"name": name, "email": email})
    except Exception as e:
        logger.debug("Zoom participant lookup failed for transcript view %s: %s", meeting_id, e)

    meeting["attendees"] = attendees
    meeting["external_attendees"] = external_attendees

    user = session.get("user", {})
    email = (user.get("preferred_username") or user.get("upn") or "").lower()
    if email:
        record_meeting_access(
            email,
            meeting.get("subject", ""),
            meeting.get("start", {}).get("dateTime", "")[:10],
        )

    return render_template(
        "transcript.html",
        meeting=meeting,
        transcript_text=transcript_text,
        attendees_json=json.dumps(attendees),
        error=None,
    )


# ── Google Meetings (email-based, no Google OAuth needed) ─────────────────────


def _extract_meeting_name(subject: str) -> str:
    """Parse meeting name from Google Meet email subject like 'Transcript for: Name'."""
    match = re.search(r"(?:transcript\s+for|recording\s+of)[:\s]+(.+)", subject, re.IGNORECASE)
    if match:
        return match.group(1).strip()
    return subject.strip()


def _normalize_google_email(msg: dict) -> dict:
    """Convert a Graph API email summary dict to the standard meeting dict."""
    received = msg.get("receivedDateTime", "")
    date_str = received[:10] if received else ""
    time_str = received[:19].replace("Z", "") if received else ""

    subject = _extract_meeting_name(msg.get("subject", "Google Meet"))

    all_recipients = msg.get("toRecipients", []) + msg.get("ccRecipients", [])
    attendees = [
        {
            "emailAddress": {
                "name": r.get("emailAddress", {}).get("name", ""),
                "address": r.get("emailAddress", {}).get("address", ""),
            }
        }
        for r in all_recipients
        if r.get("emailAddress", {}).get("address")
    ]

    return {
        "id": f"google_email_{msg['id']}",
        "subject": subject,
        "start": {"dateTime": time_str},
        "end": {"dateTime": ""},
        "organizer": {
            "emailAddress": {
                "name": msg.get("from", {}).get("emailAddress", {}).get("name", "Google Meet"),
                "address": msg.get("from", {}).get("emailAddress", {}).get("address", ""),
            }
        },
        "attendees": attendees,
        "external_attendees": [],
        "has_transcript": True,
        "source": "google",
        "received_date": date_str,
        "email_id": msg["id"],
    }


def _parse_google_meet_transcript(html_body: str) -> str:
    """
    Strip HTML from a Google Meet transcript email body and return
    a 'Speaker: text' formatted transcript string.
    """
    # Remove script/style blocks
    text = re.sub(r"<(script|style)[^>]*>.*?</\1>", "", html_body, flags=re.IGNORECASE | re.DOTALL)
    # Replace <br> tags with newlines
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.IGNORECASE)
    # Replace block-closing tags with newlines
    text = re.sub(r"</(p|div|tr|td|li|h[1-6])>", "\n", text, flags=re.IGNORECASE)
    # Strip remaining HTML tags
    text = re.sub(r"<[^>]+>", "", text)
    # Decode common HTML entities
    text = (
        text.replace("&amp;", "&")
        .replace("&lt;", "<")
        .replace("&gt;", ">")
        .replace("&nbsp;", " ")
        .replace("&#39;", "'")
        .replace("&quot;", '"')
    )
    # Collapse excessive blank lines
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]

    # Try to detect "Name: text" speaker lines and return as-is
    speaker_lines = [l for l in lines if re.match(r"^[A-Za-z][^:]{1,40}:\s+\S", l)]
    if len(speaker_lines) >= 3:
        return "\n".join(lines)

    # Fallback: join all lines as a single block
    return "\n".join(lines)


@main_bp.route("/google/dashboard")
@login_required
def google_dashboard() -> str:
    """Google Meet transcript emails dashboard — searchable by date range and keyword."""
    start_date = request.args.get("start_date", "")
    end_date = request.args.get("end_date", "")
    keyword = request.args.get("keyword", "").strip().lower()
    explicitly_searched = bool(request.args.get("start_date") or request.args.get("end_date"))

    if not start_date:
        start_date = (datetime.now() - timedelta(days=30)).strftime("%Y-%m-%d")
    if not end_date:
        end_date = datetime.now().strftime("%Y-%m-%d")

    meetings = None

    if explicitly_searched:
        token = get_token(session)
        if not token:
            flash("Your session has expired. Please sign in again.", "warning")
            return redirect(url_for("main.login"))
        try:
            gc = GraphClient(token)
            raw_emails = gc.search_google_meet_emails(start_date, end_date)
            # Filter to emails whose subject contains "transcript"
            transcript_emails = [
                e for e in raw_emails
                if "transcript" in e.get("subject", "").lower()
            ]
            # Apply optional keyword filter
            if keyword:
                transcript_emails = [
                    e for e in transcript_emails
                    if keyword in e.get("subject", "").lower()
                ]
            meetings = [_normalize_google_email(e) for e in transcript_emails]
        except Exception:
            logger.exception("google_dashboard email fetch failed")
            flash("Error fetching Google Meet transcript emails. Please try again.", "danger")
            meetings = []

    return render_template(
        "google_dashboard.html",
        meetings=meetings,
        start_date=start_date,
        end_date=end_date,
        keyword=request.args.get("keyword", ""),
    )


@main_bp.route("/google/transcript")
@login_required
def google_transcript() -> str:
    """Fetch Google Meet transcript from email body and render transcript.html."""
    user_data = session.get("user", {})
    user_email = (user_data.get("preferred_username") or user_data.get("upn") or "").lower()

    message_id = request.args.get("message_id", "")
    if not message_id:
        flash("Missing message ID.", "warning")
        return redirect(url_for("main.google_dashboard"))

    token = get_token(session)
    if not token:
        flash("Your session has expired. Please sign in again.", "warning")
        return redirect(url_for("main.login"))

    gc = GraphClient(token)
    msg = gc.get_email_message(message_id)
    if not msg:
        flash("Could not retrieve the Google Meet transcript email.", "warning")
        return redirect(url_for("main.google_dashboard"))

    meeting = _normalize_google_email(msg)

    html_body = msg.get("body", {}).get("content", "")
    transcript_text = _parse_google_meet_transcript(html_body) if html_body else ""

    error_msg = None if transcript_text.strip() else (
        "No transcript content found in the email body."
    )

    attendees = [
        {"name": a["emailAddress"]["name"], "email": a["emailAddress"]["address"]}
        for a in meeting.get("attendees", [])
    ]

    if user_email:
        record_meeting_access(
            user_email,
            meeting.get("subject", ""),
            meeting.get("received_date", ""),
        )

    return render_template(
        "transcript.html",
        meeting=meeting,
        transcript_text=transcript_text,
        attendees_json=json.dumps(attendees),
        error=error_msg,
    )
